"""Render a model-generated legal reply to a print-clean ``.docx`` file.

The model output may arrive in either of two flavours:

* **Plain prose** — what ``qwen2.5:14b`` typically returns. Each line of
  the reply maps 1:1 to a paragraph in the DOCX.
* **Markdown-flavoured** — what ``deepseek-r1:14b`` (and many
  instruction-tuned models) emits, with ``**bold**``, ``*italic*``,
  ATX headings (``# Heading``, ``## Subhead``), and ``-``/``*`` list
  bullets. If left untouched, the asterisks render as literal text in
  Word, which looks unprofessional on a legal letter.

This module normalises both flavours to a single DOCX layout:

* ``# H1`` and ``## H2`` lines  → bold paragraphs (sized 14 / 13 pt
  respectively); the leading ``#`` markers are stripped.
* ``- item`` / ``* item`` lines → bullet-list paragraphs (the marker is
  stripped).
* Inline ``**bold**`` and ``*italic*`` runs are converted to actual
  bold / italic runs in Word, with the surrounding asterisks removed.
* Blank lines are preserved as empty paragraphs to keep the visual
  spacing of the reply.

Output font is Times New Roman 12 pt, the de-facto standard for Indian
legal correspondence.
"""

from __future__ import annotations

import re
from datetime import datetime

from docx import Document
from docx.shared import Pt

from paths import OUTPUT_DIR

# Inline emphasis: ``**bold**`` first (longer match), then ``*italic*``.
# Group 1 = the inner text. Non-greedy so consecutive emphasised spans
# don't merge across each other on the same line.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)")

# A line that is just ``**Heading:**`` (or any leading/trailing whitespace
# around it) — common in deepseek-r1 output. Treated as a bold heading.
_HEADING_LINE_RE = re.compile(r"^\s*\*\*(.+?)\*\*\s*$")

# Markdown list bullet at line start: ``- ``, ``* ``, or ``+ ``. Only
# strip when followed by a space, to avoid mis-stripping ``*italic*``.
_BULLET_RE = re.compile(r"^\s*[-*+]\s+")

# ATX-style heading: one or more leading ``#`` followed by a space.
_ATX_HEADING_RE = re.compile(r"^\s*(#{1,6})\s+(.*?)\s*#*\s*$")


def save_reply_docx(reply_text: str, prefix: str = "ITax_Reply") -> Path:
    """Render ``reply_text`` to a timestamped ``.docx`` and return its path.

    :param reply_text: the model's full reply (Markdown-flavoured or
        plain prose; both are handled).
    :param prefix: filename prefix; the timestamp ``_YYYYMMDD_HHMMSS``
        is appended.
    :returns: absolute path to the written file (under
        ``backend/output/``).
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for raw_line in reply_text.splitlines():
        line = raw_line.rstrip()
        _add_line(doc, line)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"{prefix}_{timestamp}.docx"
    doc.save(out_path)
    return out_path


def _add_line(doc: Document, line: str) -> None:
    """Convert one line of Markdown-flavoured text to a DOCX paragraph."""
    if not line.strip():
        doc.add_paragraph("")
        return

    # ATX headings: ``# H1``, ``## H2``, etc.
    m = _ATX_HEADING_RE.match(line)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)
        return

    # Whole-line bold heading: ``**Subject:**``
    m = _HEADING_LINE_RE.match(line)
    if m:
        para = doc.add_paragraph()
        run = para.add_run(m.group(1))
        run.bold = True
        return

    # List bullet: ``- foo``, ``* foo``, ``+ foo``
    if _BULLET_RE.match(line):
        body = _BULLET_RE.sub("", line, count=1)
        # Use Word's built-in bullet style if available; fall back to a
        # plain paragraph with a leading bullet glyph if the style is
        # missing (it is on a default install).
        try:
            para = doc.add_paragraph(style="List Bullet")
        except KeyError:
            para = doc.add_paragraph()
            body = f"• {body}"
        _add_inline(para, body)
        return

    # Default: a normal paragraph with inline emphasis honoured.
    para = doc.add_paragraph()
    _add_inline(para, line)


def _add_inline(para, text: str) -> None:
    """Tokenise ``text`` on ``**bold**`` and ``*italic*`` and add runs.

    Bold is matched first so that ``**foo**`` cannot be mis-tokenised as
    ``*`` + ``foo`` + ``*``.
    """
    # First pass: split on bold markers. Each match becomes its own run.
    parts: list[tuple[str, str]] = []  # (kind, text)
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            parts.append(("plain", text[last : m.start()]))
        parts.append(("bold", m.group(1)))
        last = m.end()
    if last < len(text):
        parts.append(("plain", text[last:]))

    # Second pass: within each "plain" segment, split on italic markers.
    for kind, segment in parts:
        if kind == "bold":
            run = para.add_run(segment)
            run.bold = True
            continue
        last = 0
        for m in _ITALIC_RE.finditer(segment):
            if m.start() > last:
                para.add_run(segment[last : m.start()])
            run = para.add_run(m.group(1))
            run.italic = True
            last = m.end()
        if last < len(segment):
            para.add_run(segment[last:])
