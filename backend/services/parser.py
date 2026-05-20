"""Document parsers — extract plain text from notices in any supported format.

Single public entrypoint :func:`parse_file` dispatches on file extension.
All readers stream from in-memory ``bytes`` so no temporary files are written
during parsing (the upload route persists the original separately).

Supported formats and the library used:

* PDF (``.pdf``)            — :mod:`pdfplumber` (text layer only; no OCR fallback)
* Word (``.docx``)          — :mod:`docx` (paragraphs + table cells)
* Excel new (``.xlsx``)     — :mod:`openpyxl` read-only mode
* Excel legacy (``.xls``)   — :mod:`xlrd` (locked at 2.0.1 — no modern .xlsx)
* Image (``.jpg/.jpeg/.png``) — :mod:`pytesseract` + :mod:`PIL`

For image OCR, the Tesseract binary must be installed and on ``PATH``. On
Windows that is typically ``C:\\Program Files\\Tesseract-OCR\\tesseract.exe``;
if it isn't on PATH, set::

    pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

before calling :func:`parse_file`.
"""

from __future__ import annotations

import io
from datetime import date, datetime
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PIL import Image

try:
    import xlrd  # legacy .xls
except ImportError:  # pragma: no cover
    xlrd = None

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None


#: Extensions accepted by :func:`parse_file`. The upload route checks against
#: this set before reading the request body.
SUPPORTED_EXTS = {".pdf", ".docx", ".xls", ".xlsx", ".jpg", ".jpeg", ".png"}


class UnsupportedFileType(Exception):
    """Raised when :func:`parse_file` is asked to handle an unknown extension."""


def parse_file(filename: str, data: bytes) -> str:
    """Extract plain text from a notice file.

    :param filename: original filename — used only to read the extension
    :param data: raw bytes of the file
    :returns: extracted text, with sheets/tables flattened to one
        ``" | "``-separated line per row
    :raises UnsupportedFileType: if the extension isn't in :data:`SUPPORTED_EXTS`
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".xlsx":
        return _parse_xlsx(data)
    if ext == ".xls":
        return _parse_xls(data)
    if ext in (".jpg", ".jpeg", ".png"):
        return _parse_image(data)
    raise UnsupportedFileType(f"Unsupported file type: {ext}")


def _parse_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _indian_number_format(n: int | float) -> str:
    """Format an int / float with Indian-style digit grouping (1,45,000).

    Indian numbering uses a 3-digit final group then 2-digit groups from
    the right: ``1,00,000`` (one lakh), ``1,00,00,000`` (one crore). This
    helper is conservative: anything that is not a real int/float (or is
    NaN/inf) falls back to ``str(n)``. Booleans are NOT formatted as
    numbers because ``bool`` is a Python subclass of ``int``.
    """
    if isinstance(n, bool) or not isinstance(n, (int, float)):
        return str(n)
    if isinstance(n, float) and (n != n or n in (float("inf"), float("-inf"))):
        return str(n)

    sign = "-" if n < 0 else ""
    abs_n = abs(n)

    if isinstance(abs_n, float):
        whole = int(abs_n)
        frac = abs_n - whole
        # Two decimal places, dropped entirely when zero.
        if frac:
            decimal_str = f"{frac:.2f}"[1:]
            if decimal_str == ".00":
                decimal_str = ""
        else:
            decimal_str = ""
    else:
        whole = int(abs_n)
        decimal_str = ""

    whole_str = str(whole)
    if len(whole_str) <= 3:
        grouped = whole_str
    else:
        last3 = whole_str[-3:]
        rest = whole_str[:-3]
        groups: list[str] = []
        while len(rest) > 2:
            groups.insert(0, rest[-2:])
            rest = rest[:-2]
        if rest:
            groups.insert(0, rest)
        grouped = ",".join(groups) + "," + last3

    return sign + grouped + decimal_str


def _format_excel_value(value) -> str:
    """openpyxl-shaped cell value → display string.

    Rules (applied in this order — bool BEFORE int because ``isinstance(True, int)``):

    * ``None`` → empty string (placeholder preserves column alignment).
    * ``bool`` → ``"True"`` / ``"False"``.
    * ``datetime`` / ``date`` → ``DD/MM/YYYY``.
    * ``int`` / ``float`` → Indian-grouped digits.
    * Anything else → ``str(value).strip()``.
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, (int, float)):
        return _indian_number_format(value)
    return str(value).strip()


def _parse_xlsx(data: bytes) -> str:
    """Read an ``.xlsx`` and emit one ``Headers:`` line + ``Row N:`` lines per sheet.

    Column alignment is preserved by emitting ``""`` for empty / ``None``
    cells (the old implementation dropped them per-cell, collapsing column
    positions). Entirely-empty rows are still skipped — keeping them would
    just add noise without preserving any information.
    """
    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"# Sheet: {sheet}")
        header_emitted = False
        data_row_idx = 0
        for row in ws.iter_rows(values_only=True):
            cells = [_format_excel_value(c) for c in row]
            if not any(c for c in cells):
                continue
            line = " | ".join(cells)
            if not header_emitted:
                parts.append(f"Headers: {line}")
                header_emitted = True
            else:
                data_row_idx += 1
                parts.append(f"Row {data_row_idx}: {line}")
    return "\n".join(parts).strip()


def _parse_xls(data: bytes) -> str:
    """Read a legacy ``.xls`` and emit the same ``Headers:`` / ``Row N:`` format.

    xlrd does NOT auto-convert date cells to ``datetime`` — they come back
    as floats (Excel serial dates). We use the workbook's ``datemode`` and
    each cell's ``cell_type`` to recover proper types, then route through
    the same formatter as the xlsx path.
    """
    if xlrd is None:
        raise RuntimeError("xlrd is not installed; cannot parse .xls")
    book = xlrd.open_workbook(file_contents=data)
    datemode = book.datemode
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"# Sheet: {sheet.name}")
        header_emitted = False
        data_row_idx = 0
        for r in range(sheet.nrows):
            cells: list[str] = []
            for c in range(sheet.ncols):
                ct = sheet.cell_type(r, c)
                v = sheet.cell_value(r, c)
                if ct in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK, xlrd.XL_CELL_ERROR):
                    cells.append("")
                elif ct == xlrd.XL_CELL_DATE:
                    try:
                        dt = xlrd.xldate.xldate_as_datetime(v, datemode)
                        cells.append(dt.strftime("%d/%m/%Y"))
                    except Exception:
                        cells.append(str(v))
                elif ct == xlrd.XL_CELL_BOOLEAN:
                    cells.append("True" if v else "False")
                elif ct == xlrd.XL_CELL_NUMBER:
                    cells.append(_indian_number_format(v))
                else:
                    cells.append(str(v).strip())
            if not any(c for c in cells):
                continue
            line = " | ".join(cells)
            if not header_emitted:
                parts.append(f"Headers: {line}")
                header_emitted = True
            else:
                data_row_idx += 1
                parts.append(f"Row {data_row_idx}: {line}")
    return "\n".join(parts).strip()


def _parse_image(data: bytes) -> str:
    if pytesseract is None:
        raise RuntimeError("pytesseract is not installed; cannot OCR image")
    img = Image.open(io.BytesIO(data))
    # OCR: assumes Tesseract binary is on PATH. On Windows, the user may need
    # to set pytesseract.pytesseract.tesseract_cmd to the Tesseract install path.
    return pytesseract.image_to_string(img).strip()
